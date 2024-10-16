import os
import docx
import fitz
from PyQt5.QtWidgets import QApplication, QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QFileDialog, \
    QTextEdit, QListWidget, QStackedWidget, QTabWidget, QTableWidget, QTableWidgetItem, QHeaderView, QDialog, QDialogButtonBox, QSpinBox, QStyleFactory
from PyQt5.QtGui import QFont, QTextCharFormat, QColor, QTextCursor, QIcon
from PyQt5.QtCore import Qt, QThread, pyqtSignal

from spacy.tokens import Span

import spacy
import re 
import csv

from network_map_tab import NetworkMapTab
from reader_tab import ReaderTab
from editor_tab import EditorTab

class NerAnalysisThread(QThread):
    analysis_complete = pyqtSignal(dict)

    def __init__(self, document_list, label):
        super(NerAnalysisThread, self).__init__()
        self.document_list = document_list
        self.label = label

    def run(self):
        nlp = spacy.load('en_core_web_lg')
        document_entities = {}

        for document_path in self.document_list:
            try:
                if document_path.lower().endswith('.txt'):
                    with open(document_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                elif document_path.lower().endswith('.docx'):
                    doc = docx.Document(document_path)
                    content = ' '.join([paragraph.text for paragraph in doc.paragraphs])
                elif document_path.lower().endswith('.pdf'):
                    with fitz.open(document_path) as pdf_document:
                        content = ''
                        for page_number in range(pdf_document.page_count):
                            page = pdf_document[page_number]
                            content += page.get_text()

                doc = nlp(content)

                if self.label == 'SUBJECT':
                    # Combine PERSON and ORG entities into Subjects
                    entities = set(ent.text for ent in doc.ents if ent.label_ in ['PERSON', 'ORG'])
                elif self.label == 'PLACE':
                    # Combine GPE and LOC entities into Places
                    entities = set(ent.text for ent in doc.ents if ent.label_ in ['GPE', 'LOC'])
                else:
                    # Use the specified label as is
                    entities = set(ent.text for ent in doc.ents if ent.label_ == self.label)

                if document_path not in document_entities:
                    document_entities[document_path] = entities
                else:
                    entities -= document_entities[document_path]

                self.analysis_complete.emit({document_path: entities})

            except Exception as e:
                print(f"Error reading file {document_path}: {e}")



class RegexSearchThread(QThread):
    search_complete = pyqtSignal(dict)

    def __init__(self, document_list, pattern, chunk_size, overlap_size):
        super(RegexSearchThread, self).__init__()
        self.document_list = document_list
        self.pattern = pattern
        self.chunk_size = chunk_size
        self.overlap_size = overlap_size

    def run(self):
        search_results = {}

        for document_path in self.document_list:
            try:
                if document_path.lower().endswith('.txt'):
                    with open(document_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                elif document_path.lower().endswith('.docx'):
                    doc = docx.Document(document_path)
                    content = ' '.join([paragraph.text for paragraph in doc.paragraphs])
                elif document_path.lower().endswith('.pdf'):
                    with fitz.open(document_path) as pdf_document:
                        content = ''
                        for page_number in range(pdf_document.page_count):
                            page = pdf_document[page_number]
                            content += page.get_text()

                results = self.search_in_document(content)
                if results:
                    search_results[document_path] = results

            except Exception as e:
                print(f"Error reading file {document_path}: {e}")

        self.search_complete.emit(search_results)

    def search_in_document(self, content):
        results = []

        # Split the input pattern into individual terms
        search_terms = re.split(r'\s+', self.pattern.strip())

        for i in range(0, len(content) - self.chunk_size + 1, self.chunk_size - self.overlap_size):
            chunk = content[i:i + self.chunk_size]

            # Use '|'.join(search_terms) to create a regex pattern with OR operator
            pattern = re.compile(r'\b(?:' + '|'.join(map(re.escape, search_terms)) + r')\b', re.IGNORECASE)
            matches = pattern.finditer(chunk)

            for match in matches:
                start_pos = i + match.start()
                result = {
                    'position': start_pos,
                    'snippet': content[start_pos:start_pos + self.chunk_size],
                }
                results.append(result)

        return results


class DocumentReaderApp(QWidget):
    def __init__(self):
        super().__init__()

        # Set the initial directory to the directory of the script
        self.directory_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Documents")
        self.document_list = []

        self.init_ui()

        # Added variables for chunk and overlap size and maximum result size
        self.chunk_size = 200
        self.overlap_size = 20
        self.limit = 5000

        # Load documents initially
        self.load_documents()

        self.setWindowIcon(QIcon('icon.png'))

        # Apply styles
        self.apply_styles()

    def apply_styles(self):
        # Set the application style
        QApplication.setStyle(QStyleFactory.create('Fusion'))

        # Define a common font for text elements
        font = QFont()
        font.setPointSize(14)  # Adjust the font size as needed for better visibility

        # Apply styles to specific widgets
        self.setStyleSheet("""
            /* Main window background color */
            QWidget {
                background-color: #f0f0f0;
                font: bold;
            }

            /* Tab widget background color */
            QTabWidget {
                background-color: #e0e0e0;
                font: bold;
            }

            /* Sidebar background color */
            QWidget#sidebar {
                background-color: #c0c0c0;
            }

            /* Buttons */
            QPushButton {
                background-color: #418FDE;
                color: white;
                border: none;
                padding: 10px 20px;
                text-align: center;
                text-decoration: none;
                font-size: 15px;
                margin: 4px 2px;
            }

            QPushButton:hover {
                background-color: #2375c9;  /* darker color when hovered */
            }

            /* TextEdit */
            QTextEdit {
                background-color: white;
                font: """ + font.toString() + """;
            }

            /* Table */
            QTableWidget {
                background-color: white;
                font: """ + font.toString() + """;
            }
        """)

    def init_ui(self):
        # Main layout
        main_layout = QVBoxLayout()

        # Tab widget
        tab_widget = QTabWidget()

        # Tab 1
        tab1 = QWidget()
        tab1_layout = QHBoxLayout()

        # Sidebar
        sidebar_layout = QVBoxLayout()

        self.select_directory_button = QPushButton('Select Directory')
        self.select_directory_button.clicked.connect(self.select_directory)
        sidebar_layout.addWidget(self.select_directory_button)

        self.document_list_widget = QListWidget()
        self.document_list_widget.itemClicked.connect(self.show_document)
        sidebar_layout.addWidget(self.document_list_widget)

        tab1_layout.addLayout(sidebar_layout, stretch=1)  # Set stretch to 1 to make the sidebar 20% of the width

        # Content area
        content_layout = QVBoxLayout()

        self.directory_label = QLabel('Selected Directory: ')
        content_layout.addWidget(self.directory_label)

        search_controls_layout = QHBoxLayout()

        self.search_input = QTextEdit()
        self.search_input.setMaximumHeight(40)  # Set maximum height to 40 pixels
        self.search_input.setStyleSheet(
            "font-size: 14pt; line-height: 1.5; font-family: 'Calibri', monospace;")  # Set font size to 14pt, line height to 1.5, and use a monospaced font
        search_controls_layout.addWidget(self.search_input)

        self.search_button = QPushButton('Search')
        self.search_button.clicked.connect(self.search_documents)
        search_controls_layout.addWidget(self.search_button)

        content_layout.addLayout(search_controls_layout)

        self.result_label = QLabel('')
        content_layout.addWidget(self.result_label)

        self.document_viewer = QTextEdit()
        self.document_viewer.setStyleSheet(
            "font-size: 14pt; line-height: 1.5; font-family: 'Times New Roman', monospace;")  # Set font size to 12pt, line height to 1.5, and use a monospaced font
        content_layout.addWidget(self.document_viewer)

        tab1_layout.addLayout(content_layout, stretch=4)  # Set stretch to 4 to make the content area 80% of the width

        tab1.setLayout(tab1_layout)





        # Tab 2 (Empty Tab)
        tab2 = QWidget()
        tab2_layout = QVBoxLayout()

        # Buttons to choose NER label
        ner_label_buttons_layout = QHBoxLayout()
        ner_labels = ['SUBJECT', 'PLACE', 'NORP', 'FAC', 'LOC', 'PRODUCT', 'DATE', 'LAW', 'QUANTITY']  # Add more labels as needed

        self.ner_label_buttons = []
        for label in ner_labels:
            button = QPushButton(label)
            button.clicked.connect(lambda _, l=label: self.extract_entities(l))
            ner_label_buttons_layout.addWidget(button)
            self.ner_label_buttons.append(button)

        tab2_layout.addLayout(ner_label_buttons_layout)

        # Table to display results
        self.result_table = QTableWidget()
        self.result_table.setColumnCount(2)  # Two columns: Document Name and Extracted Entities
        self.result_table.setHorizontalHeaderLabels(['Document Name', 'Extracted Entities'])
        self.result_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        tab2_layout.addWidget(self.result_table)

        # Export Results Button
        export_button_layout = QHBoxLayout()
        self.export_button = QPushButton('Export Results')
        self.export_button.clicked.connect(self.export_ner_results)
        export_button_layout.addWidget(self.export_button)
        tab2_layout.addLayout(export_button_layout)

        tab2.setLayout(tab2_layout)



        # Tab 3 (Regex Search)
        tab3 = QWidget()
        tab3_layout = QVBoxLayout()

        regex_controls_layout = QHBoxLayout()

        # Button to open settings dialog
        self.settings_button = QPushButton('Settings')
        self.settings_button.clicked.connect(self.open_settings_dialog)
        regex_controls_layout.addWidget(self.settings_button)

        self.regex_input = QTextEdit()
        self.regex_input.setMaximumHeight(40)
        self.regex_input.setStyleSheet("font-size: 14pt; line-height: 1.5; font-family: 'Calibri', monospace;")
        regex_controls_layout.addWidget(self.regex_input)

        self.regex_search_button = QPushButton('Regex Search')
        self.regex_search_button.clicked.connect(self.regex_search_documents)
        regex_controls_layout.addWidget(self.regex_search_button)

        tab3_layout.addLayout(regex_controls_layout)

        # Table to display regex search results
        self.regex_result_table = QTableWidget()
        self.regex_result_table.setColumnCount(3)  # Two columns: Document Name and Matched Snippet
        self.regex_result_table.setHorizontalHeaderLabels(['Document Name', 'Matched Snippet', 'Match Count'])
        self.regex_result_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        tab3_layout.addWidget(self.regex_result_table)

        tab3.setLayout(tab3_layout)


        # Add tabs to the tab widget
        tab_widget.addTab(tab1, "Documents")
        tab_widget.addTab(tab2, "NER")
        tab_widget.addTab(tab3, "Regex Search")

        # Create an instance of NetworkMapTab and add it to the tab widget
        network_map_tab = NetworkMapTab()
        tab_widget.addTab(network_map_tab, "Network Map")

        # Create an instance of Reader Tab and add it to the tab widget
        Reader_tab = ReaderTab()
        tab_widget.addTab(Reader_tab, "Reader")

        # Create an instance of Editor Tab and add it to the tab widget
        Editor_tab = EditorTab()
        tab_widget.addTab(Editor_tab, "CSV Editor")


        main_layout.addWidget(tab_widget)
        self.setLayout(main_layout)

        self.setGeometry(300, 300, 1280, 800)
        self.setWindowTitle('Offline AI Document Analysis Tool')
        self.show()




    # Tab 3

    def regex_search_documents(self):
        if not self.directory_path:
            self.result_label.setText('Please select a directory first.')
            return

        pattern = self.regex_input.toPlainText().strip()

        if not pattern:
            self.result_label.setText('Please enter a regex pattern.')
            return

        if hasattr(self, 'regex_search_thread') and self.regex_search_thread.isRunning():
            self.result_label.setText('Regex search is already in progress.')
            return

        # Change the color and disable the regex search button
        self.regex_search_button.setStyleSheet("background-color: #A9A9A9; color: white;")
        self.regex_search_button.setEnabled(False)

        self.regex_search_thread = RegexSearchThread(self.document_list, pattern, self.chunk_size, self.overlap_size)
        self.regex_search_thread.search_complete.connect(self.display_regex_results)
        self.regex_search_thread.finished.connect(self.enable_regex_search_button)
        self.regex_search_thread.start()

    def enable_regex_search_button(self):
        # Restore the original color and enable the regex search button when the search is complete
        self.regex_search_button.setStyleSheet("")
        self.regex_search_button.setEnabled(True)


    def display_regex_results(self, result):
        self.regex_result_table.setRowCount(0)

        # Count the number of matching search terms for each result
        results_with_count = []
        for document_path, matches in result.items():
            for match in matches:
                search_text = self.regex_input.toPlainText().strip()
                if search_text:
                    search_terms = [term.strip().lower() for term in search_text.split(' ')]

                    # Count the number of matching search terms in the snippet
                    match_count = sum(term in match['snippet'].lower() for term in search_terms)
                    results_with_count.append({
                        'document_path': document_path,
                        'snippet': match['snippet'],
                        'match_count': match_count
                    })

        # Sort the results based on the number of matching search terms
        results_with_count.sort(key=lambda x: x['match_count'], reverse=True)

        # Display the sorted results in the table, up to the maximum of 5000 rows
        for result_entry in results_with_count[:self.limit]:
            row_position = self.regex_result_table.rowCount()
            self.regex_result_table.insertRow(row_position)
            self.regex_result_table.setItem(row_position, 0, QTableWidgetItem(os.path.basename(result_entry['document_path'])))
            self.regex_result_table.setItem(row_position, 1, QTableWidgetItem(result_entry['snippet'].replace('\n', ' ')))

            # Add a new column to display the number of matches
            self.regex_result_table.setItem(row_position, 2, QTableWidgetItem(str(result_entry['match_count'])))

    def open_settings_dialog(self):
        dialog = QDialog(self)
        dialog.setWindowTitle('Regex Search Settings')

        layout = QVBoxLayout()

        chunk_label = QLabel('Chunk Size:')
        chunk_spinbox = QSpinBox()
        chunk_spinbox.setValue(self.chunk_size)
        chunk_spinbox.setMinimum(20)
        chunk_spinbox.setMaximum(1500)
        chunk_spinbox.valueChanged.connect(lambda value: setattr(self, 'chunk_size', value))

        overlap_label = QLabel('Overlap Size:')
        overlap_spinbox = QSpinBox()
        overlap_spinbox.setValue(self.overlap_size)
        overlap_spinbox.setMinimum(0)
        overlap_spinbox.setMaximum(self.chunk_size - 1)
        overlap_spinbox.valueChanged.connect(lambda value: setattr(self, 'overlap_size', value))

        limit_label = QLabel('Result Limit:')
        limit_spinbox = QSpinBox()
        limit_spinbox.setValue(self.limit)
        limit_spinbox.setMinimum(1)
        limit_spinbox.setMaximum(999999)  # Adjust the maximum limit as needed
        limit_spinbox.valueChanged.connect(lambda value: setattr(self, 'limit', value))

        # Set the initial value of the spinbox based on the current limit
        limit_spinbox.setValue(self.limit)

        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)

        layout.addWidget(chunk_label)
        layout.addWidget(chunk_spinbox)
        layout.addWidget(overlap_label)
        layout.addWidget(overlap_spinbox)
        layout.addWidget(limit_label)
        layout.addWidget(limit_spinbox)
        layout.addWidget(button_box)

        dialog.setLayout(layout)

        result = dialog.exec_()
        if result == QDialog.Accepted:
            # User clicked OK, you can perform additional actions if needed
            pass
        else:
            # User clicked Cancel or closed the dialog, handle accordingly
            pass




    # Tab 2

    def extract_entities(self, label):
        self.result_table.setRowCount(0)

        # If the current thread is running, terminate it
        if hasattr(self, 'ner_analysis_thread') and self.ner_analysis_thread.isRunning():
            self.ner_analysis_thread.terminate()
            self.ner_analysis_thread.wait()

        self.ner_analysis_thread = NerAnalysisThread(self.document_list, label)
        self.ner_analysis_thread.analysis_complete.connect(self.display_entities)
        self.ner_analysis_thread.start()

    def display_entities(self, result):
        document_entities = {}
        for document_path, entities in result.items():
            document_entities[document_path] = entities

            row_position = self.result_table.rowCount()
            self.result_table.insertRow(row_position)
            self.result_table.setItem(row_position, 0, QTableWidgetItem(os.path.basename(document_path)))

            # Join entities and replace newline characters with a space
            entities_str = ', '.join(entities).replace('\n', ' ')
            self.result_table.setItem(row_position, 1, QTableWidgetItem(entities_str))

            # Use self.label instead of self.ner_label
            self.result_table.horizontalHeaderItem(1).setText(self.ner_analysis_thread.label)


    # Function to export NER results
    def export_ner_results(self):
        if self.result_table.rowCount() == 0:
            return  # No results to export

        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getSaveFileName(self, "Save NER Results", "", "CSV Files (*.csv);;All Files (*)", options=options)

        if file_name:
            with open(file_name, 'w', newline='', encoding='utf-8') as csv_file:
                csv_writer = csv.writer(csv_file)
                # Write header
                csv_writer.writerow(['Document Name', 'Extracted Entities'])
                # Write data
                for row in range(self.result_table.rowCount()):
                    document_name = self.result_table.item(row, 0).text()
                    entities = self.result_table.item(row, 1).text()
                    csv_writer.writerow([document_name, entities])






    # Tab 2


    def select_directory(self):
        directory = QFileDialog.getExistingDirectory(self, 'Select Directory')
        if directory:
            self.directory_path = directory
            self.directory_label.setText(f'Selected Directory: {directory}')
            self.load_documents()

    def load_documents(self):
        self.document_list_widget.clear()
        self.document_list = []

        for root, dirs, files in os.walk(self.directory_path):
            for file in files:
                if file.lower().endswith('.txt'):
                    self.document_list.append(os.path.join(root, file))
                    self.document_list_widget.addItem(file)
                elif file.lower().endswith('.docx'):
                    self.document_list.append(os.path.join(root, file))
                    self.document_list_widget.addItem(file)
                elif file.lower().endswith('.pdf'):
                    self.document_list.append(os.path.join(root, file))
                    self.document_list_widget.addItem(file)

    def search_documents(self):
        if not self.directory_path:
            self.result_label.setText('Please select a directory first.')
            return

        search_text = self.search_input.toPlainText().strip()

        if not search_text:
            # If no search query is entered, reload all documents into the sidebar
            self.load_documents()
            return

        search_terms = [term.strip() for term in search_text.split(' ')]
        self.update_sidebar(search_terms)

        # Clear the document viewer when search is performed
        self.document_viewer.clear()

    def update_sidebar(self, search_terms):
        self.document_list_widget.clear()

        for document_path in self.document_list:
            try:
                if document_path.lower().endswith('.txt'):
                    with open(document_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                elif document_path.lower().endswith('.docx'):
                    doc = docx.Document(document_path)
                    content = ' '.join([paragraph.text for paragraph in doc.paragraphs])
                elif document_path.lower().endswith('.pdf'):
                    with fitz.open(document_path) as pdf_document:
                        content = ''
                        for page_number in range(pdf_document.page_count):
                            page = pdf_document[page_number]
                            content += page.get_text()

                # Check if all search terms are present in the content
                if all(term.lower() in content.lower() for term in search_terms):
                    self.document_list_widget.addItem(os.path.basename(document_path))
            except Exception as e:
                print(f"Error reading file {document_path}: {e}")

    def show_document(self, item):
        document_name = item.text()
        document_path = None

        for path in self.document_list:
            if os.path.basename(path) == document_name:
                document_path = path
                break

        if document_path is not None:
            try:
                if document_path.lower().endswith('.txt'):
                    with open(document_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                elif document_path.lower().endswith('.docx'):
                    doc = docx.Document(document_path)
                    content = ' '.join([paragraph.text for paragraph in doc.paragraphs])
                elif document_path.lower().endswith('.pdf'):
                    with fitz.open(document_path) as pdf_document:
                        content = ''
                        for page_number in range(pdf_document.page_count):
                            page = pdf_document[page_number]
                            content += page.get_text()

                self.document_viewer.setPlainText(content)

                # Highlight search terms
                self.highlight_search_terms(content)

            except Exception as e:
                print(f"Error reading file {document_path}: {e}")

    def highlight_search_terms(self, content):
        cursor = self.document_viewer.textCursor()

        # Clear any previous formatting
        cursor.clearSelection()
        char_format = QTextCharFormat()
        char_format.setBackground(QColor("white"))  # Set a default background color (e.g., white)

        cursor.setCharFormat(char_format)

        search_text = self.search_input.toPlainText().strip()
        if search_text:
            search_terms = [term.strip() for term in search_text.split(' ')]

            for term in search_terms:
                cursor = self.document_viewer.textCursor()
                cursor.beginEditBlock()

                # Find all occurrences of the search term in the content (case-insensitive)
                pos = content.casefold().find(term.casefold(), 0)
                while pos != -1:
                    cursor.setPosition(pos)
                    cursor.movePosition(QTextCursor.EndOfWord, QTextCursor.KeepAnchor)
                    char_format.setBackground(QColor("yellow"))  # Highlight background color
                    cursor.setCharFormat(char_format)

                    pos = content.casefold().find(term.casefold(), pos + 1)

                cursor.endEditBlock()


if __name__ == '__main__':
    app = QApplication([])
    window = DocumentReaderApp()
    app.exec_()